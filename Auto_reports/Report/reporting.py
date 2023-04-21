import os
import requests
import openpyxl
from docx import Document
from docx.shared import Inches

# Define SharePoint credentials and URLs
USERNAME = 'your_username'
PASSWORD = 'your_password'
SHAREPOINT_URL = 'https://your_sharepoint_site.com/'
EXCEL_FILE_URL = SHAREPOINT_URL + 'path/to/excel_file.xlsx'
WORD_TEMPLATE_URL = SHAREPOINT_URL + 'path/to/word_template.docx'

# Download Excel file and Word template
response = requests.get(EXCEL_FILE_URL, auth=(USERNAME, PASSWORD))
with open('excel_file.xlsx', 'wb') as f:
    f.write(response.content)

response = requests.get(WORD_TEMPLATE_URL, auth=(USERNAME, PASSWORD))
with open('word_template.docx', 'wb') as f:
    f.write(response.content)

# Load Excel file and iterate through each sheet
workbook = openpyxl.load_workbook('excel_file.xlsx')
for sheet_name in workbook.sheetnames:
    # Create new Word document based on template
    document = Document('word_template.docx')

    # Get sheet values and inject them into Word document
    sheet = workbook[sheet_name]
    for row in sheet.iter_rows():
        for cell in row:
            for paragraph in document.paragraphs:
                if cell.coordinate in paragraph.text:
                    paragraph.text = paragraph.text.replace(cell.coordinate, str(cell.value))
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text == cell.coordinate:
                            cell.text = str(cell.value)

    # Save Word document
    document.save(sheet_name + '.docx')

# Delete downloaded files
os.remove('excel_file.xlsx')
os.remove('word_template.docx')
